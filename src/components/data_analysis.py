import os, sys, io
import seaborn as sns
import matplotlib.pyplot as plt
import pandas as pd
from dataclasses import dataclass
from src.components import data_preprocessing
from src.exception import CustomException
from src.logger import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask
from io import StringIO
import matplotlib
matplotlib.use('Agg')

app = Flask(__name__)

@dataclass
class AnalysisConfig:
    numerical_analysis_file_path = os.path.join("artifacts", "numerical.jpg")
    categorical_analysis_file_path = os.path.join("artifacts", "categorical.jpg")
    heatmap_file_path = os.path.join("artifacts", "heatmap.jpg")
    descriptive_file_path = os.path.join("artifacts", "descriptive.txt")
    zip_file_path = os.path.join("artifacts", "data_analysis.zip")
    word_doc_path = os.path.join("artifacts", "data_analysis.docx")
    X_train_path = os.path.join("artifacts","preprocessed_files","X_train.csv")

class DataAnalysis:
    def __init__(self):
        self.data_transformation_config = data_preprocessing.DataTransformationConfig()
        self.analysis_config = AnalysisConfig()

        # Ensure the directory exists for saving files
        if not os.path.exists('artifacts'):
            os.makedirs('artifacts')

        self.numerical_cols, self.categorical_cols = data_preprocessing.DataPreprocessor().NumCat()
        self.df_path = os.path.join("artifacts", "input_csv_file.csv")
        self.df = pd.read_csv(self.df_path)


        # Load numerical and categorical columns
        self.numerical_df = (
            pd.read_csv(self.data_transformation_config.numerical_cols_path)
            if os.path.exists(self.data_transformation_config.numerical_cols_path)
            else pd.DataFrame()
        )
        self.categorical_df = (
            pd.read_csv(self.data_transformation_config.categorical_cols_path)
            if os.path.exists(self.data_transformation_config.categorical_cols_path)
            else pd.DataFrame()
        )

    def create_numerical_plots(self):
        try:
            logging.info("Creating numerical plots")    
            plt.figure(figsize=(15, 15))
            plt.suptitle("Universal Analysis of Numerical Features", fontsize=20)

            num_rows = len(self.numerical_cols) // 3 + (len(self.numerical_cols) % 3 > 0)
            for i in range(len(self.numerical_cols)):
                plt.subplot(num_rows, 3, i + 1)
                sns.kdeplot(x=self.numerical_df[self.numerical_cols[i]], fill=True, color='r')  
                plt.grid(True)
                plt.xlabel(self.numerical_cols[i])

            plt.tight_layout()
            plt.savefig(self.analysis_config.numerical_analysis_file_path)
            plt.close()

            logging.info("Creating heatmap")
            plt.figure(figsize=(10, 8))
            data=pd.read_csv(self.analysis_config.X_train_path)
            sns.heatmap(data.corr(), annot=True, cmap='coolwarm')
            plt.title("Correlation Heatmap")
            plt.savefig(self.analysis_config.heatmap_file_path)
            plt.close()

        except Exception as e:
            raise CustomException(e, sys)

    def create_categorical_plots(self):
        try:
            logging.info("Creating categorical plots")
            plt.figure(figsize=(15, 15))
            plt.suptitle("Universal Analysis of Categorical Features", fontsize=20)

            n_cols = 2
            n_rows = len(self.categorical_cols) // n_cols + (len(self.categorical_cols) % n_cols > 0)

            for i in range(len(self.categorical_cols)):
                plt.subplot(n_rows, n_cols, i + 1)
                sns.countplot(x=self.categorical_df[self.categorical_cols[i]], palette='Set2')
                plt.xlabel(self.categorical_cols[i])
                plt.xticks(rotation=45)
                plt.grid(True)

            plt.tight_layout()
            logging.info("Saving categorical plot")
            plt.savefig(self.analysis_config.categorical_analysis_file_path)
            plt.close()

        except Exception as e:
            raise CustomException(e, sys)

    def descriptive_analysis(self):
        try:
            logging.info("Performing descriptive analysis")
            desc = self.df.describe()

            with open(self.analysis_config.descriptive_file_path, 'w') as f:
                f.write("Descriptive Statistics:\n")
                f.write(desc.to_string())

            buffer = StringIO()
            self.df.info(buf=buffer)
            info = buffer.getvalue()

            with open(self.analysis_config.descriptive_file_path, 'a') as f:
                f.write("\nDataFrame Info:\n")
                f.write(info)

            logging.info("Descriptive analysis saved successfully")
            
        except Exception as e:
            logging.error(f"Error during analysis: {e}")
            raise CustomException(e, sys)

    def create_word_document(self):
        try:
            logging.info("Creating Word document")
            doc = Document()
            doc.add_heading("Data Analysis Report", level=1)

            # Add Numerical Analysis
            doc.add_heading("Numerical Analysis", level=2)
            doc.add_picture(self.analysis_config.numerical_analysis_file_path, width=Inches(6))
            doc.add_paragraph("The above plot shows the distribution of numerical features.")

            # Add Heatmap
            doc.add_heading("Correlation Heatmap", level=2)
            doc.add_picture(self.analysis_config.heatmap_file_path, width=Inches(6))
            doc.add_paragraph("The above heatmap shows the correlation between  features.")

            # Add Categorical Analysis
            doc.add_heading("Categorical Analysis", level=2)
            doc.add_picture(self.analysis_config.categorical_analysis_file_path, width=Inches(6))
            doc.add_paragraph("The above plot shows the distribution of categorical features.")

            # Add Descriptive Analysis
            doc.add_heading("Descriptive Analysis", level=2)
            desc = self.df.describe()

            # Add descriptive statistics as a table
            table = doc.add_table(rows=1, cols=len(desc.columns) + 1)
            table.style = 'Table Grid'

            # Add column headers
            headers = ['Stat'] + list(desc.columns)
            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                hdr_cells[idx].text = str(header)

            # Add rows of statistics
            for index, row in desc.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)  # Add row name (e.g., count, mean, etc.)
                for idx, value in enumerate(row):
                    row_cells[idx + 1].text = f"{value:.3f}"  # Add value (formatted to 3 decimals)

            doc.add_paragraph("The above table shows the descriptive statistics for numerical features.")

            # Redirect stdout to capture df.info() output
            buffer = io.StringIO()
            sys.stdout = buffer
            self.df.info()  # Capture the output of df.info() into the buffer
            sys.stdout = sys.__stdout__  # Reset stdout

            # Get the captured info and add it to the document
            info_str = buffer.getvalue()
            doc.add_paragraph("The following information summarizes the DataFrame:")
            doc.add_paragraph(info_str)

            # Save the document
            doc.save(self.analysis_config.word_doc_path)
            logging.info(f"Word document saved: {self.analysis_config.word_doc_path}")

        except Exception as e:
            logging.error(f"Error during Word document creation: {e}")
            raise CustomException(e, sys)

    def analyze_post(self):
        try:
            self.create_numerical_plots()
            self.create_categorical_plots()
            self.descriptive_analysis()
            self.create_word_document()
            logging.info("Data analysis is completed.")
        except Exception as e:
            logging.error(f"Error during analysis: {e}")
            raise CustomException(e, sys)
